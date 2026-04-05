#!/usr/bin/env python3
"""Importador flexible de tareas: soporta CSV, XLSX y DOCX y puede usar bridge o conexión directa a Leantime.

Uso:
  python tools/import_tasks_extended.py --file tareas.xlsx --url http://localhost:8080 --token REPLACE

Si se proporciona `--token` se usa `LeantimeClient` (conexión directa a /api/jsonrpc).
Si no, se intentará POST a `--url` + `/tasks` (bridge).

Formato esperado (cabecera/tablas): `title,description,project_id`.
Para Word (`.docx`) el script busca la primera tabla con esas cabeceras.

Dependencias opcionales: `openpyxl`, `python-docx`, `httpx`.
Instalación rápida: `pip install openpyxl python-docx httpx`
"""
from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path
from typing import Any, Iterable, Iterator, List, Optional, Tuple

def read_csv(path: Path, delimiter: str = ',') -> Iterator[dict]:
    with path.open('r', encoding='utf-8', newline='') as fh:
        reader = csv.DictReader(fh, delimiter=delimiter)
        for row in reader:
            yield row

def read_xlsx(path: Path) -> Iterator[dict]:
    try:
        from openpyxl import load_workbook
    except Exception as e:
        raise RuntimeError('Falta dependencia openpyxl; instala con pip install openpyxl') from e

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    ws = wb.active
    rows = ws.iter_rows(values_only=True)
    try:
        headers = [str(h).strip() if h is not None else '' for h in next(rows)]
    except StopIteration:
        return
    for r in rows:
        obj = {headers[i]: (r[i] if i < len(r) else None) for i in range(len(headers))}
        yield {k: (str(v).strip() if v is not None else '') for k, v in obj.items()}

def read_docx(path: Path) -> Iterator[dict]:
    try:
        import docx
    except Exception as e:
        raise RuntimeError('Falta dependencia python-docx; instala con pip install python-docx') from e

    doc = docx.Document(str(path))
    # Buscar primera tabla con cabecera que incluya 'title'
    for table in doc.tables:
        # extraer primera fila como cabeceras
        hdr = [cell.text.strip() for cell in table.rows[0].cells]
        if any('title' in h.lower() for h in hdr):
            for row in table.rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                obj = {hdr[i]: values[i] if i < len(values) else '' for i in range(len(hdr))}
                yield obj
            return

    # Si no hay tabla, intentar parsear párrafos: cada párrafo = título|descripcion|project_id
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        parts = [part.strip() for part in text.split('|')]
        if not parts:
            continue
        obj = {'title': parts[0]}
        if len(parts) > 1:
            obj['description'] = parts[1]
        if len(parts) > 2:
            obj['project_id'] = parts[2]
        yield obj


def normalize_row(row: dict) -> Tuple[str, Optional[str], Optional[int]]:
    title = (row.get('title') or row.get('headline') or '').strip()
    description = (row.get('description') or row.get('desc') or '').strip() or None
    pid = row.get('project_id') or row.get('project') or row.get('projectId') or ''
    project_id = None
    if pid is not None and str(pid).strip():
        try:
            project_id = int(str(pid).strip())
        except Exception:
            project_id = None
    return title, description, project_id


def post_bridge(url: str, title: str, description: Optional[str], project_id: Optional[int]) -> Tuple[int, Any]:
    import httpx
    payload = {"title": title}
    if description:
        payload["description"] = description
    if project_id is not None:
        payload["project_id"] = project_id
    r = httpx.post(f"{url.rstrip('/')}/tasks", json=payload, timeout=30.0)
    try:
        j = r.json()
    except Exception:
        j = r.text
    return r.status_code, j


def post_direct(leantime_url: str, token: str, title: str, description: Optional[str], project_id: Optional[int]) -> dict:
    # Usa LeantimeClient local si está disponible
    try:
        from src.leantime_client import LeantimeClient
    except Exception:
        raise RuntimeError('No se puede importar LeantimeClient; asegúrate de ejecutar con PYTHONPATH=. o instala paquete')
    lc = LeantimeClient(base_url=leantime_url, api_token=token, token_header='x-api-key', token_prefix='')
    return lc.create_task(title=title, description=description, project_id=project_id)


def iter_rows_from_file(path: Path, delimiter: str = ',') -> Iterator[dict]:
    suf = path.suffix.lower()
    if suf == '.csv':
        yield from read_csv(path, delimiter=delimiter)
    elif suf in ('.xls', '.xlsx'):
        yield from read_xlsx(path)
    elif suf in ('.docx',):
        yield from read_docx(path)
    else:
        raise RuntimeError(f'Extensión no soportada: {suf}')


def main() -> None:
    p = argparse.ArgumentParser(description='Importa tareas desde CSV/XLSX/DOCX a Leantime (bridge o directo)')
    p.add_argument('--file', '-f', required=True, help='Archivo de entrada (csv/xlsx/docx)')
    p.add_argument('--url', '-u', default='http://localhost:8000', help='URL base (bridge o Leantime)')
    p.add_argument('--token', '-t', help='Token de Leantime (x-api-key). Si se proporciona, hace import directo.')
    p.add_argument('--project', '-p', type=int, help='Project ID por defecto')
    p.add_argument('--delimiter', '-d', default=',', help='Delimitador para CSV')
    args = p.parse_args()

    path = Path(args.file)
    if not path.exists():
        print('ERROR: archivo no encontrado', path, file=sys.stderr)
        sys.exit(2)

    total = 0
    created = 0
    errors = 0

    for raw in iter_rows_from_file(path, delimiter=args.delimiter):
        total += 1
        title, description, project_id = normalize_row(raw)
        if args.project:
            project_id = args.project
        if not title:
            print(f'SKIP fila {total}: falta title')
            errors += 1
            continue

        try:
            if args.token:
                res = post_direct(args.url, args.token, title, description, project_id)
                print(f'CREATED direct id={res.get("id")} title="{title}"')
            else:
                code, res = post_bridge(args.url, title, description, project_id)
                if 200 <= code < 300:
                    print(f'CREATED bridge [{code}] title="{title}" -> {res}')
                else:
                    print(f'ERROR bridge [{code}] title="{title}" -> {res}')
                    errors += 1
                    continue
            created += 1
        except Exception as e:
            print(f'ERROR fila {total}: {e}')
            errors += 1

    print(f'\nResumen: total={total} creadas={created} errores={errors}')


if __name__ == '__main__':
    main()
